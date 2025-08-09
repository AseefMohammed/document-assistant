"""
Document Processing Module
Handles ingestion and processing of various document formats (PDF, DOCX, TXT)
for the fintech document assistant RAG system.
"""

import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import hashlib
from datetime import datetime

try:
    import PyPDF2
    import docx
except ImportError:
    PyPDF2 = None
    docx = None

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Processes various document formats and splits them into chunks
    suitable for embedding and retrieval in fintech applications.
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Size of text chunks for splitting
            chunk_overlap: Overlap between chunks to maintain context
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Supported file extensions
        self.supported_extensions = {'.pdf', '.docx', '.txt', '.md'}
    
    def extract_text_from_pdf(self, file_path: str) -> tuple[str, Dict[int, str]]:
        """Extract text content from PDF file with page mapping."""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        try:
            text = ""
            page_content_map = {}
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            page_content_map[page_num + 1] = page_text
                            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
            return text.strip(), page_content_map
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text content from DOCX file."""
        if docx is None:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text content from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read().strip()
            except Exception as e:
                logger.error(f"Error reading TXT file {file_path}: {e}")
                raise
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            raise
    
    def extract_text_from_file(self, file_path: str) -> tuple[str, Dict[int, str]]:
        """
        Extract text from a file based on its extension.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (extracted text content, page content mapping)
        """
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            text = self.extract_text_from_docx(file_path)
            return text, {1: text}  # Single page for DOCX
        elif file_extension in ['.txt', '.md']:
            text = self.extract_text_from_txt(file_path)
            return text, {1: text}  # Single page for TXT/MD
        else:
            raise ValueError(f"Unsupported file extension: {file_extension}")
    
    def generate_file_hash(self, file_path: str) -> str:
        """Generate SHA-256 hash of file for change detection."""
        hash_sha256 = hashlib.sha256()
        try:
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_sha256.update(chunk)
            return hash_sha256.hexdigest()
        except Exception as e:
            logger.error(f"Error generating hash for {file_path}: {e}")
            raise
    
    def process_document(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """
        Process a document into chunks suitable for embedding.
        
        Args:
            file_path: Path to the document file
            metadata: Additional metadata to include with chunks
            
        Returns:
            List of Document objects with text chunks and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_path = os.path.abspath(file_path)
        file_name = os.path.basename(file_path)
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        logger.info(f"Processing document: {file_name}")
        
        try:
            # Extract text content and page mapping
            text_content, page_content_map = self.extract_text_from_file(file_path)
            
            if not text_content.strip():
                logger.warning(f"No text content extracted from {file_name}")
                return []
            
            # Generate file metadata
            file_stats = os.stat(file_path)
            file_metadata = {
                "source": file_path,
                "filename": file_name,
                "file_type": file_extension,
                "file_size": file_stats.st_size,
                "modified_time": datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
                "file_hash": self.generate_file_hash(file_path),
                "processed_time": datetime.now().isoformat(),
                "chunk_size": self.chunk_size,
                "chunk_overlap": self.chunk_overlap,
                "total_pages": len(page_content_map)
            }
            
            # Add custom metadata if provided
            if metadata:
                file_metadata.update(metadata)
            
            # Split text into chunks
            text_chunks = self.text_splitter.split_text(text_content)
            
            # Create Document objects
            documents = []
            for i, chunk in enumerate(text_chunks):
                # Determine which page this chunk primarily belongs to
                chunk_page = self._determine_chunk_page(chunk, page_content_map)
                
                chunk_metadata = file_metadata.copy()
                chunk_metadata.update({
                    "chunk_id": f"{file_metadata['file_hash']}_{i}",
                    "chunk_index": i,
                    "total_chunks": len(text_chunks),
                    "page_number": chunk_page,
                    "start_page": chunk_page  # For backward compatibility
                })
                
                documents.append(Document(
                    page_content=chunk,
                    metadata=chunk_metadata
                ))
            
            logger.info(f"Successfully processed {file_name} into {len(documents)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing document {file_name}: {e}")
            raise
    
    def _determine_chunk_page(self, chunk: str, page_content_map: Dict[int, str]) -> int:
        """
        Determine which page a chunk primarily belongs to based on content overlap.
        
        Args:
            chunk: The text chunk
            page_content_map: Dictionary mapping page numbers to page content
            
        Returns:
            Page number that has the highest overlap with the chunk
        """
        if not page_content_map:
            return 1
        
        max_overlap = 0
        best_page = 1
        
        # Clean the chunk for comparison (remove page markers)
        clean_chunk = chunk.replace('\n--- Page', ' Page').replace('---\n', ' ')
        
        for page_num, page_content in page_content_map.items():
            # Calculate overlap between chunk and page content
            chunk_words = set(clean_chunk.lower().split())
            page_words = set(page_content.lower().split())
            
            overlap = len(chunk_words.intersection(page_words))
            
            if overlap > max_overlap:
                max_overlap = overlap
                best_page = page_num
        
        return best_page
    
    def process_directory(self, directory_path: str, recursive: bool = True) -> List[Document]:
        """
        Process all supported documents in a directory.
        
        Args:
            directory_path: Path to the directory containing documents
            recursive: Whether to process subdirectories recursively
            
        Returns:
            List of all processed Document objects
        """
        if not os.path.exists(directory_path):
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        logger.info(f"Processing directory: {directory_path}")
        
        all_documents = []
        processed_files = 0
        
        # Get file pattern for recursive or non-recursive search
        if recursive:
            pattern = "**/*"
        else:
            pattern = "*"
        
        directory_path = Path(directory_path)
        
        for file_path in directory_path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    documents = self.process_document(str(file_path))
                    all_documents.extend(documents)
                    processed_files += 1
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    continue
        
        logger.info(f"Processed {processed_files} files into {len(all_documents)} chunks")
        return all_documents
    
    def get_document_stats(self, documents: List[Document]) -> Dict[str, Any]:
        """Get statistics about processed documents."""
        if not documents:
            return {"total_chunks": 0, "total_files": 0, "file_types": {}}
        
        stats = {
            "total_chunks": len(documents),
            "total_files": len(set(doc.metadata.get("filename", "") for doc in documents)),
            "file_types": {},
            "avg_chunk_size": 0,
            "total_text_length": 0
        }
        
        # Calculate file type distribution and text statistics
        text_lengths = []
        for doc in documents:
            file_type = doc.metadata.get("file_type", "unknown")
            stats["file_types"][file_type] = stats["file_types"].get(file_type, 0) + 1
            
            chunk_length = len(doc.page_content)
            text_lengths.append(chunk_length)
            stats["total_text_length"] += chunk_length
        
        if text_lengths:
            stats["avg_chunk_size"] = sum(text_lengths) / len(text_lengths)
        
        return stats
